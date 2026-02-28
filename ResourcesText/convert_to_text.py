#!/usr/bin/env python3
"""
Script to convert PDF and Word documents to text files.
PDF files are converted using PyPDF2.
Word files are converted using python-docx.
"""

import os
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    print("PyPDF2 not found. Install it with: pip install PyPDF2")
    exit(1)

try:
    from docx import Document
except ImportError:
    print("python-docx not found. Install it with: pip install python-docx")
    exit(1)


def convert_pdf_to_text(pdf_path, output_path):
    """Convert a PDF file to text using PyPDF2."""
    try:
        reader = PdfReader(pdf_path)
        text = []
        
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            text.append(f"=== Page {page_num} ===\n{page_text}\n")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text))
        
        print(f"✓ Converted PDF: {os.path.basename(pdf_path)} -> {os.path.basename(output_path)}")
        return True
    except Exception as e:
        print(f"✗ Error converting PDF {os.path.basename(pdf_path)}: {str(e)}")
        return False


def convert_docx_to_text(docx_path, output_path):
    """Convert a Word document to text using python-docx."""
    try:
        doc = Document(docx_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text.append('\t'.join(row_text))
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text))
        
        print(f"✓ Converted Word: {os.path.basename(docx_path)} -> {os.path.basename(output_path)}")
        return True
    except Exception as e:
        print(f"✗ Error converting Word document {os.path.basename(docx_path)}: {str(e)}")
        return False


def main():
    # Define paths
    script_dir = Path(__file__).parent
    resources_dir = script_dir.parent / 'Resources'
    output_dir = script_dir  # ResourcesText folder
    
    # Verify Resources directory exists
    if not resources_dir.exists():
        print(f"Error: Resources directory not found at {resources_dir}")
        return
    
    # Create output directory if it doesn't exist
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"Converting files from: {resources_dir}")
    print(f"Saving text files to: {output_dir}")
    print("-" * 60)
    
    # Track statistics
    stats = {'pdf': 0, 'docx': 0, 'failed': 0}
    
    # Process all files in Resources directory
    for file_path in resources_dir.iterdir():
        if file_path.is_file():
            # Get file extension
            ext = file_path.suffix.lower()
            
            # Define output file path (same name with .txt extension)
            output_path = output_dir / f"{file_path.stem}.txt"
            
            # Convert based on file type
            if ext == '.pdf':
                if convert_pdf_to_text(file_path, output_path):
                    stats['pdf'] += 1
                else:
                    stats['failed'] += 1
                    
            elif ext in ['.docx', '.doc']:
                if convert_docx_to_text(file_path, output_path):
                    stats['docx'] += 1
                else:
                    stats['failed'] += 1
            else:
                print(f"⊘ Skipped (unsupported format): {file_path.name}")
    
    # Print summary
    print("-" * 60)
    print(f"Conversion complete!")
    print(f"  PDFs converted: {stats['pdf']}")
    print(f"  Word docs converted: {stats['docx']}")
    print(f"  Failed: {stats['failed']}")


if __name__ == "__main__":
    main()
